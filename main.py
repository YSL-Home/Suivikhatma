#!/usr/bin/env python3
# ============================================================
#   DocBot — @DocpoursalmaBot
#   Image / PDF  →  Word (.docx)
#   100% Python — aucune dépendance Node.js
# ============================================================

import os, json, base64, logging, tempfile, re
from pathlib import Path
from datetime import datetime
from io import BytesIO

import anthropic
import fitz                         # PyMuPDF  — PDF → images
from docx import Document           # python-docx — génération .docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from telegram import Update
from telegram.ext import (
    Application, CommandHandler, MessageHandler,
    filters, ContextTypes
)

# ─────────────────────────────────────────────────────────────
#  CONFIG
# ─────────────────────────────────────────────────────────────

TELEGRAM_TOKEN = os.environ.get(
    "TELEGRAM_TOKEN",
    "8276662498:AAE4HI-i3ymz9dliHC9_8oVGU1FOZ0dANfE"
)
CLAUDE_API_KEY = os.environ.get(
    "CLAUDE_API_KEY",
    "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"
)

WORK_DIR = Path(tempfile.gettempdir()) / "docbot"
WORK_DIR.mkdir(exist_ok=True)

logging.basicConfig(
    format="%(asctime)s [%(levelname)s] %(message)s",
    level=logging.INFO
)
log = logging.getLogger(__name__)

# ─────────────────────────────────────────────────────────────
#  PROMPT CLAUDE
# ─────────────────────────────────────────────────────────────

SYSTEM_PROMPT = """
Tu es un expert en retranscription fidèle de documents médicaux/scientifiques manuscrits ou imprimés.

MISSION : Extraire TOUT le contenu de l'image et retourner un JSON structuré pour recréer le document Word à l'identique.

RÈGLES ABSOLUES :
1. Copie EXACTEMENT chaque mot, chiffre, symbole visible — ne résume pas, ne reformule pas.
2. Respecte la hiérarchie visuelle : titre principal > titres de sections > sous-titres > texte.
3. Conserve les flèches (→, ←, ↓, ↑), symboles (≥, ≤, +, ×, °), abréviations médicales.
4. Si tu vois un encadré ou boîte de texte → type "box".
5. Si tu vois un schéma/algorithme/arbre décisionnel → type "flowchart_node" avec flèches "flowchart_arrow".
6. Ignore les couleurs de surlignage — document final tout en noir.
7. Annotations manuscrites dans les marges → type "annotation".
8. Conserve le gras et souligné quand clairement visibles.
9. Chaque ligne distincte = un bloc séparé.

TYPES DE BLOCS :
- "title"           : titre principal (centré, grand)
- "heading1"        : titre niveau 1 (ex: I. INTRODUCTION)
- "heading2"        : titre niveau 2 (ex: A. Sous-section)
- "heading3"        : titre niveau 3
- "paragraph"       : texte normal
- "bullet"          : puce niveau 1
- "bullet2"         : puce niveau 2
- "bullet3"         : puce niveau 3
- "numbered"        : liste numérotée
- "arrow_point"     : ligne commençant par → ou + ou *
- "box"             : texte encadré (NB, ATTENTION, définition)
- "box_bullet"      : puce dans un encadré
- "flowchart_node"  : nœud d'algorithme (rectangle/ovale)
- "flowchart_arrow" : flèche de connexion dans schéma
- "annotation"      : note manuscrite dans marge
- "separator"       : ligne vide
- "table_header"    : en-tête de tableau (colonnes séparées par " | ")
- "table_row"       : ligne de tableau (colonnes séparées par " | ")

FORMAT JSON — UNIQUEMENT ce JSON, rien d'autre, sans backticks :
{
  "title": "Titre principal",
  "blocks": [
    {"type": "heading1", "text": "I. INTRODUCTION", "bold": true, "underline": true, "italic": false},
    {"type": "bullet",   "text": "Contenu...",       "bold": false, "underline": false, "italic": false}
  ]
}
"""

# ─────────────────────────────────────────────────────────────
#  ANALYSE IMAGE → JSON  (Claude Vision)
# ─────────────────────────────────────────────────────────────

def prepare_image(image_path: str) -> str:
    """
    Prépare l'image pour l'API Claude :
    - Compresse si > 4MB (limite Anthropic ~5MB)
    - Upscale si trop petite (< 800px)
    - Convertit en JPEG
    Retourne le chemin de l'image prête.
    """
    from PIL import Image as PILImage
    
    MAX_BYTES = 4 * 1024 * 1024   # 4 MB max
    MIN_WIDTH = 800                # px minimum
    
    img  = PILImage.open(image_path).convert("RGB")
    w, h = img.size
    size = os.path.getsize(image_path)
    log.info("Image originale : %dx%d px, %d KB", w, h, size // 1024)
    
    out  = str(WORK_DIR / ("prep_%s.jpg" % Path(image_path).stem))
    qual = 90
    
    # Upscale si trop petite
    if w < MIN_WIDTH:
        factor = MIN_WIDTH / w
        img = img.resize((int(w * factor), int(h * factor)), PILImage.LANCZOS)
        w, h = img.size
        log.info("Upscalée → %dx%d", w, h)
    
    # Sauvegarder et vérifier la taille
    img.save(out, "JPEG", quality=qual)
    
    # Réduire qualité jusqu'à < 4MB
    while os.path.getsize(out) > MAX_BYTES and qual > 40:
        qual -= 10
        img.save(out, "JPEG", quality=qual)
        log.info("Compression qualité=%d → %d KB", qual, os.path.getsize(out)//1024)
    
    # Si encore trop grand, réduire la résolution
    if os.path.getsize(out) > MAX_BYTES:
        scale = 0.75
        while os.path.getsize(out) > MAX_BYTES and scale > 0.3:
            nw, nh = int(w * scale), int(h * scale)
            img.resize((nw, nh), PILImage.LANCZOS).save(out, "JPEG", quality=60)
            log.info("Redim → %dx%d, %d KB", nw, nh, os.path.getsize(out)//1024)
            scale -= 0.15
    
    final_size = os.path.getsize(out)
    log.info("Image prête : %d KB (qualité=%d)", final_size//1024, qual)
    return out


def analyze_image(image_path: str) -> dict:
    size_bytes = os.path.getsize(image_path)
    log.info("Analyse Claude : %s (%d bytes)", image_path, size_bytes)

    # Préparer l'image (compression + qualité optimale)
    image_path = prepare_image(image_path)

    client = anthropic.Anthropic(api_key=CLAUDE_API_KEY)
    ext  = Path(image_path).suffix.lower()
    mime = {".jpg":"image/jpeg",".jpeg":"image/jpeg",
            ".png":"image/png",".webp":"image/webp"}.get(ext, "image/jpeg")

    with open(image_path, "rb") as f:
        img_b64 = base64.standard_b64encode(f.read()).decode()

    log.info("Envoi à Claude : %d chars base64, mime=%s", len(img_b64), mime)

    def _call_claude(prompt_text):
        return client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=8000,
            system=SYSTEM_PROMPT,
            messages=[{
                "role": "user",
                "content": [
                    {"type": "image",
                     "source": {"type": "base64", "media_type": mime, "data": img_b64}},
                    {"type": "text", "text": prompt_text}
                ]
            }]
        )

    def _parse_json(raw):
        raw = raw.strip()
        if "```" in raw:
            for part in raw.split("```"):
                p = part.strip()
                if p.startswith("json"): p = p[4:].strip()
                if p.startswith("{"): raw = p; break
        start = raw.find("{")
        end   = raw.rfind("}") + 1
        if start >= 0 and end > start:
            return json.loads(raw[start:end])
        return json.loads(raw)

    # ── Tentative 1 : prompt standard ──────────────────────
    prompt1 = (
        "Analyse cette image de notes médicales/scientifiques et extrait TOUT le texte visible.\n"
        "Retourne UNIQUEMENT ce JSON (sans backticks, sans texte avant ou après) :\n"
        '{"title": "titre principal du document", "blocks": [\n'
        '  {"type": "heading1", "text": "I. INTRODUCTION", "bold": true, "underline": true, "italic": false},\n'
        '  {"type": "paragraph", "text": "texte normal...", "bold": false, "underline": false, "italic": false},\n'
        '  {"type": "bullet", "text": "item de liste", "bold": false, "underline": false, "italic": false}\n'
        "]}"
    )

    try:
        msg  = _call_claude(prompt1)
        raw  = msg.content[0].text
        log.info("Réponse Claude T1 (300 chars): %s", raw[:300])
        data = _parse_json(raw)
        nb   = len(data.get("blocks", []))
        log.info("T1 → %d blocs, titre=\'%s\'", nb, data.get("title", ""))
        if nb > 0:
            return data
        log.warning("T1 : 0 blocs, tentative 2...")
    except Exception as e:
        log.error("T1 erreur : %s", e)

    # ── Tentative 2 : prompt simplifié sans JSON schema ────
    prompt2 = (
        "Lis TOUT le texte de cette image ligne par ligne.\n"
        "Retourne un JSON avec :\n"
        "- title : le titre principal\n"
        "- blocks : liste de {type, text, bold, underline, italic}\n"
        "type = paragraph pour tout le texte (ne cherche pas à catégoriser).\n"
        "JSON UNIQUEMENT, aucun autre texte."
    )

    try:
        msg  = _call_claude(prompt2)
        raw  = msg.content[0].text
        log.info("Réponse Claude T2 (300 chars): %s", raw[:300])
        data = _parse_json(raw)
        nb   = len(data.get("blocks", []))
        log.info("T2 → %d blocs", nb)
        if nb > 0:
            return data
        log.warning("T2 : toujours 0 blocs")
    except Exception as e:
        log.error("T2 erreur : %s", e)

    # ── Tentative 3 : demander du texte brut puis convertir ──
    prompt3 = (
        "Transcris TOUT le texte visible dans cette image, ligne par ligne.\n"
        "Réponds UNIQUEMENT avec le texte transcrit, rien d'autre."
    )

    try:
        msg      = _call_claude(prompt3)
        raw_text = msg.content[0].text.strip()
        log.info("T3 texte brut (300 chars): %s", raw_text[:300])

        if len(raw_text) > 10:
            # Convertir le texte brut en blocs
            lines  = [l.strip() for l in raw_text.split("\n") if l.strip()]
            title  = lines[0] if lines else "Document"
            blocks = []
            for line in lines[1:]:
                btype = "heading1" if line.startswith(("I.", "II.", "III.", "IV.", "V.")) else                         "bullet"   if line.startswith(("•", "-", "*", "+")) else                         "arrow_point" if line.startswith("→") else                         "paragraph"
                blocks.append({"type": btype, "text": line,
                                "bold": False, "underline": False, "italic": False})
            log.info("T3 → %d blocs construits depuis texte brut", len(blocks))
            return {"title": title, "blocks": blocks}
    except Exception as e:
        log.error("T3 erreur : %s", e)

    # Aucune tentative n'a fonctionné
    return {"title": "Document", "blocks": []}


def pdf_to_images(pdf_path: str, ts: str) -> list:
    doc = fitz.open(pdf_path)
    paths = []
    for i, page in enumerate(doc):
        pix  = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0))
        path = str(WORK_DIR / ("pdf_p%d_%s.jpg" % (i, ts)))
        pix.save(path)
        paths.append(path)
    doc.close()
    log.info("PDF → %d image(s)", len(paths))
    return paths


def merge_pages(pages: list) -> dict:
    if len(pages) == 1:
        return pages[0]
    blocks = []
    title  = pages[0].get("title", "Document")
    for i, p in enumerate(pages):
        if i > 0:
            blocks.append({"type":"separator","text":"","bold":False,"underline":False,"italic":False})
        blocks.extend(p.get("blocks", []))
    return {"title": title, "blocks": blocks}

# ─────────────────────────────────────────────────────────────
#  GÉNÉRATION .DOCX  (python-docx — 100% Python)
# ─────────────────────────────────────────────────────────────

def _set_cell_border(cell):
    """Ajoute une bordure noire à toutes les faces d'une cellule."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement('w:' + side)
        border.set(qn('w:val'),   'single')
        border.set(qn('w:sz'),    '4')
        border.set(qn('w:color'), '000000')
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _add_run(para, text: str, bold=False, underline=False, italic=False, size_pt=11):
    run = para.add_run(text)
    run.bold      = bold
    run.underline = underline
    run.italic    = italic
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = 'Arial'
    return run


def _para_border(doc_obj, text, bold=False, italic=False, indent=False):
    """Paragraphe entouré d'une bordure noire (encadré)."""
    para = doc_obj.add_paragraph()
    pPr  = para._p.get_or_add_pPr()

    pBdr = OxmlElement('w:pBdr')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement('w:' + side)
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    '4')
        el.set(qn('w:color'), '000000')
        el.set(qn('w:space'), '4')
        pBdr.append(el)
    pPr.append(pBdr)

    if indent:
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), '720')
        pPr.append(ind)

    _add_run(para, text, bold=bold, italic=italic)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(4)
    return para


def generate_docx(data: dict, out_path: str) -> str:
    doc   = Document()
    style = doc.styles['Normal']
    style.font.name  = 'Arial'
    style.font.size  = Pt(11)
    style.font.color.rgb = RGBColor(0, 0, 0)

    # Marges
    for section in doc.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(0.75)
        section.right_margin  = Inches(0.75)

    title_text = data.get("title", "Document")
    blocks     = data.get("blocks", [])

    # ── Titre principal ──────────────────────────────────────
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = _add_run(tp, title_text, bold=True, underline=True, size_pt=16)
    tp.paragraph_format.space_after = Pt(12)

    # ── Blocs ────────────────────────────────────────────────
    for b in blocks:
        t     = b.get("type", "paragraph")
        text  = b.get("text", "")
        bold  = b.get("bold",      False)
        uline = b.get("underline", False)
        ital  = b.get("italic",    False)

        if not text and t not in ("separator",):
            continue

        # ── Séparateur ──
        if t == "separator":
            doc.add_paragraph()

        # ── Titres ──
        elif t == "title":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_run(p, text, bold=True, underline=True, size_pt=14)

        elif t == "heading1":
            p = doc.add_heading(level=1)
            p.clear()
            _add_run(p, text, bold=True, underline=True, size_pt=13)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)

        elif t == "heading2":
            p = doc.add_heading(level=2)
            p.clear()
            _add_run(p, text, bold=True, underline=True, size_pt=12)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)

        elif t == "heading3":
            p = doc.add_heading(level=3)
            p.clear()
            _add_run(p, text, bold=True, size_pt=11)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)

        # ── Listes à puces ──
        elif t == "bullet":
            p = doc.add_paragraph(style='List Bullet')
            _add_run(p, text, bold=bold, underline=uline, italic=ital)

        elif t == "bullet2":
            p = doc.add_paragraph(style='List Bullet 2')
            _add_run(p, text, bold=bold, italic=ital)

        elif t == "bullet3":
            p = doc.add_paragraph(style='List Bullet 3')
            _add_run(p, text, bold=bold, italic=ital)

        # ── Liste numérotée ──
        elif t == "numbered":
            p = doc.add_paragraph(style='List Number')
            _add_run(p, text, bold=bold, italic=ital)

        # ── Flèche / point ──
        elif t == "arrow_point":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            _add_run(p, text, bold=bold, italic=ital)

        # ── Encadrés ──
        elif t in ("box", "box_bullet"):
            _para_border(doc, text, bold=bold, italic=ital,
                         indent=(t == "box_bullet"))

        # ── Schéma / flowchart ──
        elif t == "flowchart_node":
            _para_border(doc, text, bold=bold)
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        elif t == "flowchart_arrow":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            arrow_text = "↓  " + text if text else "↓"
            _add_run(p, arrow_text, italic=True, size_pt=10)

        # ── Annotation ──
        elif t == "annotation":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(3.0)
            _add_run(p, "[%s]" % text, italic=True, size_pt=9)

        # ── Tableau ──
        elif t in ("table_header", "table_row"):
            cols   = text.split(" | ")
            ncols  = max(len(cols), 1)
            is_hdr = (t == "table_header")
            table  = doc.add_table(rows=1, cols=ncols)
            table.style = 'Table Grid'
            row = table.rows[0]
            for j, col_text in enumerate(cols):
                cell = row.cells[j]
                cell.text = ""
                p = cell.paragraphs[0]
                _add_run(p, col_text.strip(), bold=is_hdr)
                _set_cell_border(cell)

        # ── Paragraphe normal ──
        else:
            p = doc.add_paragraph()
            _add_run(p, text, bold=bold, underline=uline, italic=ital)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)

    doc.save(out_path)
    log.info("DOCX généré : %s (%d bytes)", out_path, os.path.getsize(out_path))
    return out_path

# ─────────────────────────────────────────────────────────────
#  HANDLERS TELEGRAM
# ─────────────────────────────────────────────────────────────

async def cmd_start(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        "👋 Bonjour ! Je suis *DocBot* 📄\n\n"
        "📸 Envoie une *photo* ou *image* (JPG/PNG)\n"
        "📄 Ou un *PDF* (toutes les pages → 1 seul .docx)\n\n"
        "→ Je te renvoie un *.docx* propre, tout en noir\n\n"
        "💡 Envoie en *fichier* (📎) pour meilleure qualité\n"
        "⏱ Traitement : 30 à 90 secondes",
        parse_mode="Markdown"
    )


async def cmd_help(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        "📖 *Aide — DocBot*\n\n"
        "✅ *Formats acceptés :*\n"
        "• 🖼️ Image : JPG, PNG, WEBP\n"
        "• 📄 PDF : toutes les pages → 1 seul .docx\n\n"
        "✅ *Ce qui est préservé :*\n"
        "• Titres et sous-titres\n"
        "• Listes à puces / numérotées\n"
        "• Encadrés (NB, ATTENTION...)\n"
        "• Schémas et algorithmes\n"
        "• Gras, souligné, italique\n"
        "• Flèches et symboles médicaux\n\n"
        "🚫 Surlignages couleur → tout en noir\n\n"
        "💡 Envoie *en fichier* (📎) pour éviter la compression",
        parse_mode="Markdown"
    )


async def process_input(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    msg = update.message
    log.info("Fichier de %s (id=%s)", msg.from_user.first_name, msg.from_user.id)

    ts        = datetime.now().strftime("%Y%m%d_%H%M%S%f")
    docx_path = str(WORK_DIR / ("doc_%s.docx" % ts))
    tmp_files = []

    # ── Détecter le type ──
    is_pdf = is_image = False
    if msg.photo:
        is_image = True; label = "image"
    elif msg.document:
        mime = msg.document.mime_type or ""
        if   mime == "application/pdf":      is_pdf = True;  label = "PDF"
        elif mime.startswith("image/"):      is_image = True; label = "image"
        else:
            await msg.reply_text(
                "❌ Format non supporté.\n\n"
                "✅ Acceptés :\n• Images : JPG, PNG, WEBP\n• PDF")
            return
    else:
        return

    wait = await msg.reply_text(
        "⏳ *%s reçu !* Traitement en cours..." % label.upper(),
        parse_mode="Markdown"
    )

    try:
        # ── 1. Télécharger ──
        await wait.edit_text("📥 *1/3* — Téléchargement...", parse_mode="Markdown")

        if msg.photo:
            fobj = await ctx.bot.get_file(
                sorted(msg.photo, key=lambda p: p.file_size, reverse=True)[0].file_id)
            dl = str(WORK_DIR / ("img_%s.jpg" % ts))
            await fobj.download_to_drive(dl)
            tmp_files.append(dl)
            image_paths = [dl]

        elif is_image:
            fobj = await ctx.bot.get_file(msg.document.file_id)
            dl = str(WORK_DIR / ("img_%s.jpg" % ts))
            await fobj.download_to_drive(dl)
            tmp_files.append(dl)
            image_paths = [dl]

        else:  # PDF
            fobj = await ctx.bot.get_file(msg.document.file_id)
            pdf_path = str(WORK_DIR / ("pdf_%s.pdf" % ts))
            await fobj.download_to_drive(pdf_path)
            tmp_files.append(pdf_path)

            await wait.edit_text(
                "🔄 *1/3* — Conversion PDF en images...",
                parse_mode="Markdown")
            image_paths = pdf_to_images(pdf_path, ts)
            tmp_files.extend(image_paths)

            nb_p = len(image_paths)
            if nb_p > 15:
                await wait.edit_text(
                    "⚠️ PDF de %d pages — traitement des 15 premières." % nb_p,
                    parse_mode="Markdown")
                image_paths = image_paths[:15]

        nb_pages = len(image_paths)

        # ── 2. Analyser ──
        await wait.edit_text(
            "🤖 *2/3* — Analyse IA (%d page(s))...\n_Extraction texte et structure_" % nb_pages,
            parse_mode="Markdown")

        pages_data = []
        for i, img_path in enumerate(image_paths):
            if nb_pages > 1:
                await wait.edit_text(
                    "🤖 *2/3* — Analyse page %d / %d..." % (i+1, nb_pages),
                    parse_mode="Markdown")
            try:
                pages_data.append(analyze_image(img_path))
            except Exception as e:
                log.warning("Page %d erreur : %s", i+1, e)
                pages_data.append({"title": "Page %d" % (i+1), "blocks": []})

        doc_data = merge_pages(pages_data)
        nb_blocs = len(doc_data.get("blocks", []))
        titre    = doc_data.get("title", "Document")

        # Avertir si aucun contenu extrait
        if nb_blocs == 0:
            await wait.edit_text(
                "⚠️ Aucun contenu extrait de l'image.\n\n"
                "Conseils :\n"
                "• Envoie l'image *en fichier* (📎) pour éviter la compression\n"
                "• Assure-toi que l'image est bien éclairée et nette\n"
                "• Pour un PDF, envoie-le directement en fichier",
                parse_mode="Markdown"
            )
            return

        # ── 3. Générer .docx ──
        await wait.edit_text(
            "📄 *3/3* — Génération du Word...\n_%d éléments extraits_" % nb_blocs,
            parse_mode="Markdown")
        generate_docx(doc_data, docx_path)
        tmp_files.append(docx_path)

        # ── 4. Envoyer ──
        safe  = re.sub(r"[^\w\s\-]", "", titre)[:45].strip() or "document"
        fname = "%s.docx" % safe
        pages_info = " · %d pages" % nb_pages if nb_pages > 1 else ""

        with open(docx_path, "rb") as f:
            await msg.reply_document(
                document=f, filename=fname,
                caption=(
                    "✅ *%s*\n\n"
                    "📊 %d éléments extraits%s\n"
                    "🖤 Mise en forme propre — tout en noir\n"
                    "📐 Structure originale préservée"
                ) % (fname, nb_blocs, pages_info),
                parse_mode="Markdown"
            )
        await wait.delete()

    except json.JSONDecodeError:
        await wait.edit_text(
            "⚠️ Contenu difficile à lire.\n"
            "Essaie avec une image plus nette ou un PDF de meilleure qualité.",
            parse_mode="Markdown")
    except Exception as e:
        log.error("Erreur: %s", e, exc_info=True)
        await wait.edit_text(
            "❌ Erreur :\n`%s`\n\nRéessaie." % str(e)[:150],
            parse_mode="Markdown")
    finally:
        for p in tmp_files:
            try:
                if os.path.exists(p): os.unlink(p)
            except Exception:
                pass


async def msg_other(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        "📸 Envoie une *image* ou un *PDF* → je génère un Word.\n/help pour l'aide.",
        parse_mode="Markdown")

# ─────────────────────────────────────────────────────────────
#  MAIN
# ─────────────────────────────────────────────────────────────

def main():
    log.info("DocBot démarré — @DocpoursalmaBot (python-docx, no Node.js)")

    app = Application.builder().token(TELEGRAM_TOKEN).build()
    app.add_handler(CommandHandler("start", cmd_start))
    app.add_handler(CommandHandler("help",  cmd_help))
    app.add_handler(MessageHandler(filters.PHOTO,                                process_input))
    app.add_handler(MessageHandler(filters.Document.IMAGE,                       process_input))
    app.add_handler(MessageHandler(filters.Document.MimeType("application/pdf"), process_input))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND,              msg_other))
    app.run_polling(drop_pending_updates=True)


if __name__ == "__main__":
    main()
